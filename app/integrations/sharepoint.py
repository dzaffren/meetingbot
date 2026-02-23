from __future__ import annotations

import io
import logging
from datetime import datetime

from docx import Document
from docx.shared import Pt
from kiota_abstractions.base_request_configuration import RequestConfiguration

from app.config import get_settings
from app.integrations.graph_client import get_graph_client
from app.models.minutes import MeetingMinutes

logger = logging.getLogger(__name__)


def _render_docx(minutes: MeetingMinutes) -> bytes:
    """Render MeetingMinutes to a .docx byte buffer."""
    doc = Document()

    doc.add_heading(minutes.title, level=1)
    doc.add_paragraph(f"Date: {minutes.date.isoformat()}    Meeting ID: {minutes.meeting_id}")

    doc.add_heading("Attendees", level=2)
    doc.add_paragraph(", ".join(minutes.attendees) if minutes.attendees else "N/A")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(minutes.summary)

    if minutes.key_decisions:
        doc.add_heading("Key Decisions", level=2)
        for decision in minutes.key_decisions:
            doc.add_paragraph(f"• {decision}")

    if minutes.action_items:
        doc.add_heading("Action Items", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Task"
        hdr[1].text = "PIC"
        hdr[2].text = "Due Date"
        hdr[3].text = "Planner Task ID"
        for item in minutes.action_items:
            row = table.add_row().cells
            row[0].text = item.title
            row[1].text = item.pic
            row[2].text = item.due_date.isoformat() if item.due_date else ""
            row[3].text = item.planner_task_id or ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def upload_minutes(minutes: MeetingMinutes) -> str:
    """
    Render MeetingMinutes to .docx and upload to SharePoint.
    Returns the SharePoint web URL of the uploaded file.
    """
    settings = get_settings()
    graph = get_graph_client()

    filename = f"{minutes.date.isoformat()} - {minutes.title}.docx"
    docx_bytes = _render_docx(minutes)

    # Graph API: PUT /sites/{site-id}/drives/{drive-id}/items/{folder-path}:/{filename}:/content
    upload_path = f"{settings.sharepoint_folder_path}/{filename}"

    try:
        result = await (
            graph.sites.by_site_id(settings.sharepoint_site_id)
            .drives.by_drive_id(settings.sharepoint_drive_id)
            .root.item_with_path(upload_path)
            .content.put(docx_bytes)
        )
        url = result.web_url or ""
        logger.info("Uploaded minutes to SharePoint: %s", url)
        return url
    except Exception as exc:
        logger.error("Failed to upload minutes to SharePoint: %s", exc)
        raise
